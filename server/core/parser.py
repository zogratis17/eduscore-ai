import fitz  # PyMuPDF
import docx
import io
from typing import Dict, Any

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extracts text from PDF bytes using PyMuPDF.
        
        """
        try:
            # Open the PDF from bytes
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            if doc.is_encrypted:
                return {"error": "The PDF is encrypted and cannot be processed."}
            
            full_text = ""
            
            for page in doc:
                full_text += page.get_text()
                
            return {
                "status": "success",
                "text": full_text.strip(),
                "metadata": {
                    "page_count": len(doc),
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", "")
                }
            }
        
        except Exception as e:
            return {"error":"error","message": f"An error occurred while processing the PDF: {str(e)}"}
        
    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> Dict[str, Any]:
        
        try:
            
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)
            
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                    
            final_text = "\n".join(full_text).strip()
            
            return {
                "status": "success",
                "text": final_text.strip(),
                "metadata": {
                    "paragraph_count": len(doc.paragraphs),
                    "author": doc.core_properties.author or "unknown"
                }
            }
            
        except Exception as e:
            return {"error":"error","message": f"An error occurred while processing the DOCX: {str(e)}"}
       
        
        
            
        
        pass
    
        