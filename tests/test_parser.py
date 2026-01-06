import pytest
import io
import docx
from server.core.parser import DocumentParser

def create_dummy_docx() -> bytes:
    """Helper to create a DOCX in memory"""
    doc = docx.Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")
    doc.core_properties.author = "Test Bot"
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def test_extract_text_from_pdf():
    file_path = "tests/sample.pdf"
    
    try:
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        
        result = DocumentParser.extract_text_from_pdf(file_bytes)
        
        assert result["status"] == "success"
        assert "text" in result
        assert isinstance(result["text"], str)
        assert "metadata" in result
        assert "page_count" in result["metadata"]
        assert result["metadata"]["page_count"] > 0
        
        # Preview the text in the console
        print(f"\n--- Extracted Text (Total: {len(result['text'])} chars) ---")
        print("FIRST 300 CHARS:")
        print(result["text"][:300])
        print("\n... [MIDDLE CONTENT] ...\n")
        print("LAST 300 CHARS:")
        print(result["text"][-300:])
        print("--------------------------------------------------")
    except FileNotFoundError:
        pytest.fail(f"Test file not found at {file_path}")
    except Exception as e:
        pytest.fail(f"An unexpected error occurred: {e}")

def test_extract_text_from_docx():
    # 1. Create a fake DOCX in memory
    docx_bytes = create_dummy_docx()
    
    # 2. Extract
    result = DocumentParser.extract_text_from_docx(docx_bytes)
    
    # 3. Assert
    assert result["status"] == "success"
    assert "Hello World" in result["text"]
    assert "This is a test document." in result["text"]
    assert result["metadata"]["paragraph_count"] >= 2
    assert result["metadata"]["author"] == "Test Bot"

def test_extract_text_from_real_docx():
    file_path = "tests/sample.docx"
    try:
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        
        result = DocumentParser.extract_text_from_docx(file_bytes)
        
        assert result["status"] == "success"
        print(f"\n--- REAL DOCX PREVIEW ---")
        print(result["text"][:300])
        print("--------------------------")
    except FileNotFoundError:
        pytest.skip(f"No real sample.docx found at {file_path}, skipping.")